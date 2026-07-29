import os
import json
import base64
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime

import httpx
from pypdf import PdfReader

from app.config import settings
from app.services.llm_service import llm_service

logger = logging.getLogger(__name__)

MISTRAL_OCR_URL = "https://api.mistral.ai/v1/ocr"

TEXT_EXTENSIONS = {".txt", ".md"}
OCR_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg"}   # sent straight to Mistral OCR
DOCX_EXTENSIONS = {".docx"}                          # parsed locally, no API call needed
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | OCR_EXTENSIONS | DOCX_EXTENSIONS

MIME_BY_EXT = {
    ".pdf": "application/pdf",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
}


class ResumeService:
    def __init__(self):
        self.resumes_dir = os.path.join("app", "data", "resumes")
        self.jd_path = os.path.join("app", "data", "job_description.json")

    DEFAULT_CAREEM_JD = {
        "title": "Senior Backend Engineer (Python) - Ride Matching & Dispatch Team",
        "company": "Careem",
        "department": "Ride Hailing Core",
        "location": "Dubai, UAE (Remote Friendly)",
        "description": "Careem is looking for a Senior Backend Engineer to join our Core Matching & Dispatch Team. In this role, you will design, build, and optimize high-throughput, low-latency backend systems and APIs that power our real-time vehicle matching, dynamic surge pricing, and driver routing algorithms. You will work on distributed state machines processing millions of concurrent bookings daily, ensuring system availability, sub-second latency, and horizontal scalability.",
        "requirements": [
            "5+ years of production experience designing and building scalable backend systems in Python.",
            "Strong proficiency in modern Python frameworks such as FastAPI, Django, or Flask.",
            "Deep understanding of microservices architecture, event-driven design, and API patterns (REST, gRPC, WebSockets).",
            "Hands-on experience with relational databases (PostgreSQL), data modeling, and caching mechanisms (Redis).",
            "Experience with message brokers and event streaming networks (Apache Kafka, RabbitMQ).",
            "Familiarity with cloud infrastructures (AWS), Docker, and Kubernetes deployment workflows.",
            "Proven track record of optimizing systems for high concurrency, distributed locking, and microsecond latencies."
        ],
        "preferred_qualifications": [
            "Prior experience working in ride-hailing, delivery, logistics, or last-mile transit companies.",
            "Knowledge of Go or Java for cross-service development and integration.",
            "Familiarity with geospatial index libraries (Uber H3, Google S2) or PostGIS for location-based query optimization."
        ]
    }

    # ---------------------------------------------------------------- JD ----

    def load_job_description(self) -> Dict[str, Any]:
        if not os.path.exists(self.jd_path):
            self.save_job_description(self.DEFAULT_CAREEM_JD)
            return self.DEFAULT_CAREEM_JD
        with open(self.jd_path, "r", encoding="utf-8") as f:
            return json.load(f)

    def save_job_description(self, jd_data: Dict[str, Any]) -> Dict[str, Any]:
        os.makedirs(os.path.dirname(self.jd_path), exist_ok=True)
        with open(self.jd_path, "w", encoding="utf-8") as f:
            json.dump(jd_data, f, indent=2)
        return jd_data

    def reset_job_description(self) -> Dict[str, Any]:
        return self.save_job_description(self.DEFAULT_CAREEM_JD)

    # ---------------------------------------------------------- Conversion ----

    def convert_upload_to_markdown(self, file_bytes: bytes, filename: str) -> str:
        """
        Converts an uploaded resume (PDF, DOCX, PNG, JPG, TXT or MD) into clean Markdown.
        PDFs and images go straight to Mistral's OCR API (one fast call, no local
        document-parsing library). DOCX is parsed locally with python-docx (it already
        has a text layer, so no OCR/LLM call is needed). TXT/MD pass through as-is.
        """
        ext = os.path.splitext(filename)[1].lower()

        if ext in TEXT_EXTENSIONS:
            return file_bytes.decode("utf-8", errors="ignore").strip()

        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format '{ext}'. Please upload PDF, DOCX, PNG, JPG, TXT or MD.")

        if ext in DOCX_EXTENSIONS:
            markdown_content = self._parse_docx(file_bytes)
            if not markdown_content:
                raise ValueError("No text could be extracted from this DOCX file. It may be empty or corrupted.")
            return markdown_content

        # ext in OCR_EXTENSIONS (.pdf, .png, .jpg, .jpeg)
        if not settings.MISTRAL_API_KEY:
            if ext == ".pdf":
                # No key configured -- fall back to a local, text-layer-only PDF read.
                markdown_content = self._parse_pdf_fallback(file_bytes)
                if markdown_content:
                    return markdown_content
            raise ValueError(
                "No MISTRAL_API_KEY configured, so this file can't be OCR'd. "
                "Add MISTRAL_API_KEY to your .env, or upload a DOCX/TXT/MD instead."
            )

        try:
            markdown_content = self._mistral_ocr(file_bytes, ext)
        except Exception as e:
            logger.error(f"Mistral OCR failed for {filename}: {e}")
            if ext == ".pdf":
                markdown_content = self._parse_pdf_fallback(file_bytes)
                if markdown_content:
                    return markdown_content
            raise ValueError(f"Could not read this file via Mistral OCR: {e}") from e

        if not markdown_content:
            raise ValueError("No text could be extracted from this file. It may be empty, blank, or unreadable.")

        return markdown_content

    def _mistral_ocr(self, file_bytes: bytes, ext: str) -> str:
        """Sends a PDF or image straight to Mistral's OCR API (mistral-ocr-latest) and
        returns the combined Markdown for all pages. One HTTP call, no local parsing library."""
        mime_type = MIME_BY_EXT[ext]
        b64_data = base64.b64encode(file_bytes).decode("utf-8")
        data_uri = f"data:{mime_type};base64,{b64_data}"

        document_payload = (
            {"type": "image_url", "image_url": data_uri}
            if ext in {".png", ".jpg", ".jpeg"}
            else {"type": "document_url", "document_url": data_uri}
        )

        response = httpx.post(
            MISTRAL_OCR_URL,
            headers={
                "Authorization": f"Bearer {settings.MISTRAL_API_KEY}",
                "Content-Type": "application/json",
            },
            json={"model": settings.MISTRAL_OCR_MODEL, "document": document_payload},
            timeout=60.0,
        )
        response.raise_for_status()
        data = response.json()

        pages = data.get("pages", [])
        return "\n\n".join(page.get("markdown", "") for page in pages).strip()

    def _parse_docx(self, file_bytes: bytes) -> str:
        """Extracts headings/paragraphs/tables from a DOCX into simple Markdown."""
        from io import BytesIO
        from docx import Document

        doc = Document(BytesIO(file_bytes))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower()
            if "heading 1" in style or "title" in style:
                lines.append(f"# {text}")
            elif "heading" in style:
                lines.append(f"## {text}")
            elif "list" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))

        return "\n".join(lines).strip()

    def _parse_pdf_fallback(self, file_bytes: bytes) -> str:
        """Last-resort raw text extraction if MarkItDown returns nothing for a PDF."""
        try:
            from io import BytesIO
            reader = PdfReader(BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Fallback PDF parsing also failed: {e}")
            return ""

    # ------------------------------------------------------------ Resumes ----

    def list_resumes(self) -> List[Dict[str, str]]:
        resumes = []
        if not os.path.exists(self.resumes_dir):
            return resumes

        for file_name in sorted(os.listdir(self.resumes_dir)):
            if file_name.endswith((".md", ".txt")):
                candidate_id = os.path.splitext(file_name)[0]
                name = candidate_id.replace("_", " ").title()
                resumes.append({
                    "id": candidate_id,
                    "name": name,
                    "file_name": file_name,
                    "file_type": "markdown"
                })
        return resumes

    def get_resume_content(self, candidate_id: str) -> str:
        for file_name in os.listdir(self.resumes_dir):
            if os.path.splitext(file_name)[0] == candidate_id:
                file_path = os.path.join(self.resumes_dir, file_name)
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
        raise FileNotFoundError(f"Candidate resume not found for ID: {candidate_id}")

    # ------------------------------------------------------------ Scoring ----

    def screen_candidate(self, candidate_id: str) -> Dict[str, Any]:
        job_desc = self.load_job_description()
        resume_text = self.get_resume_content(candidate_id)

        evaluation = llm_service.screen_resume(job_desc, resume_text)
        evaluation["candidate_id"] = candidate_id
        evaluation["candidate_name"] = candidate_id.replace("_", " ").title()
        evaluation["evaluation_date"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        return evaluation

    def screen_custom_resume(self, file_name: str, file_bytes: bytes) -> Dict[str, Any]:
        """Converts an uploaded resume to Markdown, saves it, and screens it against the active JD."""
        markdown_content = self.convert_upload_to_markdown(file_bytes, file_name)

        job_desc = self.load_job_description()
        evaluation = llm_service.screen_resume(job_desc, markdown_content)

        candidate_name = llm_service._infer_candidate_name(markdown_content)
        base_name = os.path.splitext(file_name)[0].lower().replace(" ", "_")
        candidate_id = "custom_" + "".join(c for c in base_name if c.isalnum() or c in ("_", "-"))

        try:
            os.makedirs(self.resumes_dir, exist_ok=True)
            custom_file_path = os.path.join(self.resumes_dir, f"{candidate_id}.md")
            with open(custom_file_path, "w", encoding="utf-8") as f:
                f.write(markdown_content)
        except Exception as e:
            logger.error(f"Failed to save custom resume {candidate_id} to disk: {e}")

        evaluation["candidate_id"] = candidate_id
        evaluation["candidate_name"] = candidate_name
        evaluation["evaluation_date"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        return evaluation

    # ----------------------------------------------------------- Feedback ----

    def get_resume_feedback(self, candidate_id: str) -> Dict[str, Any]:
        """Generates resume-improvement suggestions for an already-loaded candidate."""
        job_desc = self.load_job_description()
        resume_text = self.get_resume_content(candidate_id)
        feedback = llm_service.generate_resume_feedback(job_desc, resume_text)
        feedback["candidate_id"] = candidate_id
        return feedback


resume_service = ResumeService()
